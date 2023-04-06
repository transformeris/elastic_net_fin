import pandas as pd
import pickle
import matplotlib.font_manager as fm
import numpy as np
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches
from matplotlib import pyplot as plt
import pandas as pd
import akshare as ak
def load_obj(name):
    with open(name + '.pkl', 'rb') as f:
        return pickle.load(f)

def get_max_drawdown(data):
    cum_return = data['cumulative_return']
    drawdown = (cum_return / cum_return.cummax()) - 1
    max_drawdown = drawdown.min()

    end_date = drawdown.idxmin()
    start_date = cum_return[:end_date].idxmax()

    return max_drawdown, start_date, end_date

def plot_cumulative_return(data):
    plt.figure(figsize=(10, 6))
    font = fm.FontProperties(fname='C:\Windows\Fonts/msyh.ttc')
    plt.plot(data['cumulative_return'], label='策略累计收益')
    plt.plot(data['benchmark_cumulative_return'], label='基准累计收益')
    plt.legend(prop=font)
    plt.title('策略与基准累计收益', fontproperties=font)
    plt.savefig('cumulative_return.png', dpi=300)
def backtest_strategy_T1(data, benchmark, commission, slippage, start_date, end_date):
    data = data[(data.index >= start_date) & (data.index <= end_date)]
    benchmark = benchmark[(benchmark.index >= start_date) & (benchmark.index <= end_date)]

    data = data.loc[start_date:end_date]
    benchmark = benchmark.loc[start_date:end_date]

    # 初始化现金和持仓
    cash = 100
    position = 0

    # 计算每日收益
    data['daily_return'] = (data['close'] - data['open'] - slippage) / (data['open'] + slippage) - commission

    # 初始化策略收益和基准收益
    data['strategy_return'] = 0
    data['benchmark_return'] = (benchmark['close'] / benchmark['close'].shift(1)) - 1
    data['cash'] = 0
    data['position'] = 0
    data['last_trade_date'] = pd.NaT # 最后一次交易日期

    for i in range(1, len(data)):
        # 判断是否满足T+1交易限制
        if not pd.isnull(data.iloc[i - 1]['last_trade_date']):
            days_since_last_trade = (data.iloc[i]['date'] - data.iloc[i - 1]['last_trade_date']).days
            if days_since_last_trade <= 1:
                data.iloc[i, data.columns.get_loc('strategy_return')] = 0
                data.iloc[i, data.columns.get_loc('cash')] = cash
                data.iloc[i, data.columns.get_loc('position')] = position
                data.iloc[i, data.columns.get_loc('last_trade_date')] = data.iloc[i - 1]['last_trade_date']
                continue

        # 买入
        buy_value = min(cash, data.iloc[i]['open'])
        shares_to_buy = buy_value / (data.iloc[i]['open'] + slippage)
        cash -= shares_to_buy * (data.iloc[i]['open'] + slippage)

        # 增加持仓
        position += shares_to_buy

        # 卖出
        sell_value = min(position * data.iloc[i]['close'], position * (data.iloc[i]['close'] - slippage))
        shares_to_sell = sell_value / (data.iloc[i]['close'] - slippage)
        cash += shares_to_sell * (data.iloc[i]['close'] - slippage)

        # 减少持仓
        position -= shares_to_sell

        # 更新最后交易日期
        data.iloc[i, data.columns.get_loc('last_trade_date')] = data.iloc[i]['date']

        # 计算策略收益
        data.iloc[i, data.columns.get_loc('strategy_return')] = (sell_value - buy_value) / (buy_value + slippage)
        data.iloc[i, data.columns.get_loc('cash')] = cash
        data.iloc[i, data.columns.get_loc('position')] = position

    return data

def backtest_strategy(data, benchmark, commission, slippage, start_date, end_date):
    data = data[(data.index >= start_date) & (data.index <= end_date)]
    benchmark = benchmark[(benchmark.index >= start_date) & (benchmark.index <= end_date)]

    data = data.loc[start_date:end_date]
    benchmark = benchmark.loc[start_date:end_date]

    # 初始化现金和持仓
    cash = 100
    position = 0

    # 计算每日收益
    data['daily_return'] = (data['close'] - data['open'] - slippage) / (data['open'] + slippage) - commission

    # 初始化策略收益和基准收益
    data['strategy_return'] = 0
    data['benchmark_return'] = (benchmark['close'] / benchmark['close'].shift(1)) - 1
    data['cash'] = 0
    data['position'] = 0
    for i in range(1, len(data)):
        # 买入
        buy_value = min(cash, data.iloc[i]['open'])
        shares_to_buy = buy_value / (data.iloc[i]['open'] + slippage)
        cash -= shares_to_buy * (data.iloc[i]['open'] + slippage)

        # 增加持仓
        position += shares_to_buy

        # 卖出
        sell_value = min(position * data.iloc[i]['close'], position * (data.iloc[i]['close'] - slippage))
        shares_to_sell = sell_value / (data.iloc[i]['close'] - slippage)
        cash += shares_to_sell * (data.iloc[i]['close'] - slippage)

        # 减少持仓
        position -= shares_to_sell

        # 计算策略收益
        data.iloc[i, data.columns.get_loc('strategy_return')] = (sell_value - buy_value) / (buy_value + slippage)
        data.iloc[i, data.columns.get_loc('cash')] = cash
        data.iloc[i, data.columns.get_loc('position')] = position
    # data['daily_return'] = (data['close'] - data['open']) / data['open'] * (1 - commission - slippage)
    # data['daily_return'] = (data['close']- data['open'].shift(1)) / data['open'].shift(1)* (1 - commission - slippage)
    data['cumulative_return'] = (1 + data['strategy_return']).cumprod()
    data['benchmark_daily_return'] = (benchmark['close'] - benchmark['close'].shift(1)) / benchmark['close'].shift(1)
    data['benchmark_cumulative_return'] = (1 + data['benchmark_daily_return']).cumprod()
    data['excess_return'] = data['cumulative_return'] - data['benchmark_cumulative_return']



    # data['drawdown'] = (data['cumulative_return'] / data['cumulative_return'].cummax()) - 1
    # max_drawdown = data['drawdown'].min()
    # max_drawdown_period = data[data['drawdown'] == max_drawdown].index.tolist()
    max_drawdown, max_drawdown_start, max_drawdown_end = get_max_drawdown(data)
    annualized_return = data['cumulative_return'][-1] ** (252 / len(data)) - 1
    annualized_benchmark_return = data['benchmark_cumulative_return'][-1] ** (252 / len(data)) - 1
    annualized_excess_return = annualized_return - annualized_benchmark_return
    data['annualized_return']=annualized_return
    annualized_volatility = data['daily_return'].std() * np.sqrt(252)
    annualized_benchmark_volatility = data['benchmark_daily_return'].std() * np.sqrt(252)
    sharpe_ratio = (annualized_return - 0.02) / annualized_volatility

    alpha = annualized_excess_return

    plot_cumulative_return(data)

    doc = Document()
    doc.add_heading('回测报告', 0)
    doc.add_heading('策略表现', level=1)
    doc.add_paragraph(f'年化收益: {data["annualized_return"][-1]:.4f}')
    doc.add_paragraph(f'策略累计收益: {data["cumulative_return"][-1]:.4f}')
    doc.add_paragraph(f'基准累计收益: {data["benchmark_cumulative_return"][-1]:.4f}')
    doc.add_paragraph(f'策略最大回撤: {max_drawdown:.4f}')
    doc.add_paragraph(f'最大回撤发生时间段: {max_drawdown_start} 至 {max_drawdown_end}')
    doc.add_paragraph(f'阿尔法收益: {alpha:.4f}')
    doc.add_paragraph(f'夏普比率: {sharpe_ratio:.4f}')
    doc.add_heading('图形展示', level=1)
    doc.add_picture('cumulative_return.png', width=Inches(6))

    doc.save('回测报告.docx')

    return data
if __name__ == '__main__':
    # data = pd.read_csv('data.csv', index_col=0, parse_dates=True)
    # benchmark = pd.read_csv('benchmark.csv', index_col=0, parse_dates=True)
    etf_kline_all = load_obj('etf_all')
    zhengquan_kline = etf_kline_all['sz159919']
    zhengquan_kline.loc[:, 'trade_date'] = zhengquan_kline.index
    pd.to_datetime(zhengquan_kline.loc[:, 'trade_date'])
    zhengquan_kline.set_index(pd.to_datetime(zhengquan_kline.loc[:, 'trade_date']), inplace=True)
    data_proto=zhengquan_kline
    data = zhengquan_kline
    benchmark = zhengquan_kline

    commission = 0.00006
    slippage = 0
    start_date = '2014-01-01'
    end_date = '2016-01-01'

    result = backtest_strategy(data, benchmark, commission, slippage, start_date, end_date)
    print(result)