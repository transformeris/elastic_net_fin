import docx

doc=docx.Document(r'D:\腾讯\微信\WeChat Files\wxid_xv36th150smv12\FileStorage\File\2022-08\执照考证题库.docx')
# docx.opendocx('C:/Users\梁海镇\Desktop\执照考证题库.docx')
zz=doc.paragraphs[120].text
res={}
res2={}
for num in range(600):#range(len(doc.paragraphs)):
    # print(doc.paragraphs[num].text)
    if '.' in doc.paragraphs[num].text:
        timu=doc.paragraphs[num].text.split('.')
        try:
            float(timu[0])
            res[num]=timu
#
        except:
            pass
#     if '答案' in doc.paragraphs[num].text:
#
#         res2[num]=doc.paragraphs[num].text[-2]
#
#     if doc.paragraphs[num].text=='':
#         print('fuck')
#
#         doc.paragraphs[num]._element.remove
#         p=doc.paragraphs[num]
#         # p.getparent().remove(p)
#
# for para in doc.paragraphs:
#     if para.text=='':
#         para.clear()
#        # p=para._element
#        # p.get_parent().remove(p)
