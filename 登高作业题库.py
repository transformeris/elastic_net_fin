import docx
import re
doc = docx.Document('D:\elink下载\WXWorkLocal\File/2023-08\登高作业600题(含解析）.docx')
res=[]
import pandas as pd
doc_text = ''

# Loop through each paragraph in the document
for para in doc.paragraphs:
    # Add the text of the paragraph to the document text
    doc_text += para.text + '\n'

res1 = []

# Initialize a variable to keep track of the question number
question_num = 0

# Loop through each paragraph in the document
for para in doc.paragraphs:
    # Check if the paragraph starts with a number followed by a period
    if para.text.startswith('1.') or para.text.startswith('2.') or para.text.startswith('3.') or para.text.startswith('4.') or para.text.startswith('5.'):
        # Increment the question number
        question_num += 1

        # Add the text of the question to the list
        res1.append(para.text.strip() + ' ')

    # Check if the paragraph starts with "答案：" or "A " or "B " or "C "
    elif para.text.startswith('答案：') or para.text.startswith('A ') or para.text.startswith('B ') or para.text.startswith('C '):
        # Add the text of the answer to the last question in the list
        res1[-1] += para.text.strip() + ' '

# Print the list of questions and answers
print(res1)


# Initialize a list to store the questions and answers
res = []

# Initialize a variable to keep track of the question number
question_num = 0

# Loop through each paragraph in the document
for para in doc.paragraphs:
    # Check if the paragraph starts with a number followed by a period
    if para.text.startswith(tuple([f"{i}." for i in range(1, 1000)])):
        # Increment the question number
        question_num += 1

        # Add the text of the question to the list
        res.append(para.text.strip() + ' ')

    # Check if the paragraph starts with "答案：" or "A " or "B " or "C "
    elif para.text.startswith('答案：') or para.text.startswith('A') or para.text.startswith('B') or para.text.startswith('C'):
        # Add the text of the answer to the last question in the list
        res[-1] += para.text.strip() + ' '

# Print the list of questions and answers
print(res)




patternA = r'A(.*?)B'
patternB = r'B(.*?)C'
patternC = r'C(.*?)$'

pattern = r'（(A|B|C)）'
timu=[]
daan_A=[]
daan_B=[]
daan_C=[]
res_final={}
num=0
res2={}
xuanzheti=[]
panduanti=[]
for i in res:
    if "A" in i and "B" in i and "C" in i:
        xuanzheti.append(i)
    else:
        panduanti.append(i)


pattern = r'（(A|B|C|D)）'

res_xuan={}
num=0
for i in xuanzheti:
    num=num+1
    res_xuan2={}
    print(i)
    daan= re.search(pattern, i).group(1)

    i = re.sub(pattern, '（）', i)
    res_xuan2['题目']=i[i.find('.')+1:i.find('A')]
    res_xuan[num]=res_xuan2
    if "D" not in i or 'ZLD' in i or 'LED' in i:
        res_xuan2['A']=re.search(patternA, i).group(1)
        res_xuan2['B']=re.search(patternB, i).group(1)
        res_xuan2['C']=re.search(patternC, i).group(1)
    elif "D" in i and 'ZLD' not in i and 'LED' not in i:
        res_xuan2['A']=re.search(patternA, i).group(1)
        res_xuan2['B']=re.search(patternB, i).group(1)
        res_xuan2['C']=re.search(r'C(.*?)D', i).group(1)
        res_xuan2['D']=re.search(r'D(.*?)$', i).group(1)
    res_xuan2['答案'] = daan

# pd.DataFrame(res_xuan).T.to_excel('D:\新建文件夹 (3)\OneDrive - 7x541z\桌面\登高作业600题(含解析）2.xlsx')
# pd.Series(panduanti).to_excel('D:\新建文件夹 (3)\OneDrive - 7x541z\桌面\登高作业600题(含解析）.xlsx')